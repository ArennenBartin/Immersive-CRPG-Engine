from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "AI_GAME_BUILD_GUIDE_FOR_CODEX.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)


def set_table_borders(table, color="B8C4D2", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_width(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, True)
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_callout(doc: Document, title: str, body: str, fill: str = "F2F7FC") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="8AA9C7")
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=180, bottom=120, end=180)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(4)


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(4)


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("1F4D78")
    title.paragraph_format.space_after = Pt(4)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build_doc() -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("AI Game Build Guide for the CRPG Engine")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "A Codex-oriented authoring manual for building games, content, maps, events, assets, and systemic encounters in this engine."
    ).italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Project: CRPG Engine Base Toolkit | Updated: 2026-07-05 | Format: compact reference guide")

    add_callout(
        doc,
        "Purpose",
        "This document is written for AI agents and human collaborators using Codex to build a game inside the current engine. It explains what the engine is, how the editor is organized, what data tools exist, and how to use switches, conditions, triggers, cutscenes, maps, assets, quests, simulation, and audits without inventing a second engine.",
        fill="EAF4FF",
    )

    add_heading(doc, "1. What This Engine Is", 1)
    add_para(
        doc,
        "The project is a React, TypeScript, Vite, Zustand, Zod, and Tailwind browser-based CRPG engine and editor. It has two faces: a Studio/editor for authoring data and a Play runtime for testing the game. The current active game view is a 2D grid renderer, GameRenderer2D. The engine is not a free-form scene graph; it is a package-driven grid CRPG toolkit where the playable world is described by structured data and the runtime resolves that data into movement, fog, combat, dialogue, quests, simulation, and persistent save deltas.",
    )
    add_para(
        doc,
        "The core spatial contract is a single x/z grid. One tile is one cell and one step. Cells can be walkable or blocked, can block line of sight, can hold terrain/surface/simulation tags, and can belong to rooms or regions. Oblique or painted art is a render skin over that square grid. It does not change collision, pathing, line of sight, fog, chemistry, exits, or trigger cells.",
    )
    add_para(
        doc,
        "The game content lives in a package with schema literal crpg_engine_game_package_v1. Saves track runtime state, switches, fog memory, map deltas, chemistry, combat, inventory, faction reputation, Alderamontico emotional state, and other persistent consequences. When Codex builds content, it should usually add or modify package data and use existing runtime systems before changing engine code.",
    )

    add_table(
        doc,
        ["Layer", "What It Owns", "AI Authoring Rule"],
        [
            ["Studio/editor", "Map, object, sprite, dialogue, quest, event, item, shop, skill, document, and simulation authoring.", "Use the editor and schemas as the source of truth for content shape."],
            ["Game package", "The data model: maps, objects, sprites, entities, dialogue, quests, cutscenes, switches, items, skills, shops, barks, and simulation records.", "Prefer data additions over bespoke runtime branches."],
            ["Play runtime", "Movement, fog, interaction, combat, events, barks, shops, inventory, documents, chemistry, Grid/Attend, AI, and save deltas.", "Test changes in Play mode and preserve persistent state rules."],
            ["Engine core", "Deterministic helpers, command/effect/event dispatch, story predicates, combat, chemistry, simulation, immersive-sim systems, and adapters.", "Only edit when a feature cannot be expressed in package data or existing verbs."],
            ["Audits/tests", "Map reachability, combat validity, chemistry/state behavior, overworld asset/map contracts, TypeScript validation, and builds.", "Every nontrivial slice ends with audits, lint, build, or targeted tests."],
        ],
        widths=[1.35, 2.9, 2.25],
    )

    add_heading(doc, "2. What Codex and AI Can Do Here", 1)
    add_para(
        doc,
        "Codex is most useful as a builder, auditor, and continuity keeper. It can read the docs, inspect schemas, edit content and code, generate art prompts or assets, run extraction scripts, use the browser to visually test the editor, and write updated documentation. The best workflow is to keep each change as a vertical slice: data, runtime behavior, UI feedback, audit, and documentation all move together.",
    )
    add_table(
        doc,
        ["AI Tool", "Use It For", "Good Output"],
        [
            ["Repository search", "Find schemas, stores, renderer paths, runtime commands, and existing patterns.", "A grounded implementation plan using real file names and field names."],
            ["Code editing", "Patch TypeScript, scripts, docs, schemas, presets, generators, and data registries.", "Small scoped diffs that follow existing patterns."],
            ["Terminal scripts", "Run npm scripts, extraction scripts, audits, tests, TypeScript lint, and production builds.", "A passing command set or a precise failure report."],
            ["In-app browser QA", "Open localhost, test Play mode, inspect visual output, verify overlays, and reproduce editor/runtime issues.", "Screenshots or observations tied to a concrete bug or acceptance pass."],
            ["Image generation", "Create new terrain, object, prop, player, and NPC source art when the style/perspective contract is known.", "Source atlases that can be cropped, normalized, edge-conditioned, and registered."],
            ["Asset scripts", "Crop atlases, remove gutters/backgrounds, write manifests/contact sheets, and wire sprite ids.", "Deterministic asset files plus auditable manifests."],
            ["Document generation", "Produce build guides, phase summaries, design specs, audits, and handoff docs.", "Rendered and visually checked DOCX/PDF artifacts."],
        ],
        widths=[1.55, 2.55, 2.4],
    )

    add_heading(doc, "3. The AI Build Loop", 1)
    add_numbers(
        doc,
        [
            "Read the current docs stack and the relevant schema/code before deciding how to build the slice.",
            "Identify whether the requested feature is content, UI, runtime behavior, asset pipeline, or audit coverage.",
            "Make the smallest real vertical slice: authoring data, renderer/play feedback, save persistence, and tests where needed.",
            "Use existing package fields, conditions, cutscenes, triggers, global verbs, simulation records, and scripts before adding a new system.",
            "Run the relevant command set: npm run lint, npm run build, npm run audit:maps, npm run audit:combat, npm run audit:overworld, npm run audit:overworld-assets, npm run test:chemistry, npm run test:state, or npm run test:engine.",
            "Open the browser for visual QA whenever the change touches maps, Play mode, HUD, sprites, overlays, or editor controls.",
            "Update docs or manifests when a phase, data contract, art standard, or authoring rule changes.",
        ],
    )
    add_callout(
        doc,
        "AI guardrail",
        "Do not invent parallel gameplay layers. The removed Praxis/casework layer is intentionally gone. Interpretation, story state, and player-facing choices should route through the normal dialogue, quest, document, event, switch, Attend, and simulation systems.",
        fill="FFF4E5",
    )

    add_heading(doc, "4. Editor Interface Map", 1)
    add_para(
        doc,
        "The left sidebar is the main control surface. It switches between Home, Map, Play, Tiles, Sprites, Dialogue, Quests, Entities, Events, Items, Documents, Shops, Skills, and Simulation. The top bar provides undo/redo and the current editor name. Play mode hides extra chrome on mobile so the runtime has more screen space.",
    )
    add_table(
        doc,
        ["Editor", "Purpose", "Use When Building"],
        [
            ["Home", "Project entry point and general package workflow.", "Orient, import/export, or return to the package overview."],
            ["Map", "Paint cells, place objects/entities/triggers/exits/items/containers, inspect map data, run generation/settlement helpers, lint, fit view, and jump to Play.", "Build spaces, traversal, zones, set-pieces, map exits, and trigger cells."],
            ["Play", "Runtime test surface for movement, fog, combat, interaction, dialogue, shops, inventory, events, simulation, Attend, and UI feedback.", "Validate the game as a player, not just as data."],
            ["Tiles", "Author object/tile records and pixel/data-url tile sprites bound by tile_sprite_id.", "Create or revise placeable floors, walls, doors, and object tiles."],
            ["Sprites", "Create and manage sprite records, including generated or pixel sprites.", "Wire player/NPC/object art and confirm sprite ids."],
            ["Dialogue", "Build dialogue trees: nodes, speakers, text, scene images, options, quest/switch/cutscene effects, and option gates.", "Write NPC scenes, choices, story state changes, and short interactions."],
            ["Quests", "Define quests and objectives: talk, kill, collect, explore, interact, or custom.", "Track player goals and objective completion hooks."],
            ["Entities", "Define actors: stats, sprite, dialogue, NPC flag, skills, emotional axes, Attend nodes.", "Make NPCs, enemies, party members, and authored Attend targets."],
            ["Events", "Author cutscenes/action sequences used by triggers and dialogue options.", "Script state changes, teleports, combat starts, rewards, shops, music, fades, and endings."],
            ["Items", "Define inventory items: category, icon/sprite, effects, spatial profile, simulation profile.", "Make healing items, keys, equipment, supplies, and systemic props."],
            ["Documents", "Write readable in-game documents.", "Create notes, letters, journals, terminal text, and lore fragments."],
            ["Shops", "Define shop inventory, base prices, conditional stock, and conditional price modifiers.", "Build merchants, barter gates, and reputation-based prices."],
            ["Skills", "Define AP/MP costs, element, targeting shape, range, payloads, and emotional impulses.", "Create combat, support, elemental, and emotional verbs."],
            ["Simulation", "Inspect and author material, object/Part, tile-layer, scheduler, reaction, perception, global-verb, combat, inventory, region, workstation, and process surfaces.", "Tune systemic behavior and verify simulation records."],
        ],
        widths=[1.15, 2.45, 2.9],
    )

    add_heading(doc, "5. Map Editor Operating Notes", 1)
    add_para(
        doc,
        "The Map editor is where game spaces become real. The toolbar includes a map selector, width/height controls, Y/layer controls, Generate, Settlement, validation/lint tools, Inspector, Play map, and Fit. The bottom tool rail contains the main authoring tools.",
    )
    add_table(
        doc,
        ["Map Tool", "What It Writes", "Authoring Notes"],
        [
            ["Walkable", "Sets cells active/walkable.", "Use for floors, roads, and traversable terrain. Brush size applies."],
            ["Wall", "Sets cells blocked.", "Use for void, walls, closed barriers, and impassable terrain. Remember LOS separately through blocks_los and object/cell rules."],
            ["Raise / Lower", "Adjusts height/visual height.", "Use for terrain readability and height/combat experiments; keep traversal audits passing."],
            ["Spawn", "Adds or moves map spawns with facing.", "Every map needs a reliable start or edge spawn for exits."],
            ["Tile", "Assigns terrain/object tile fields to cells.", "Use broad cohesive patches, not random per-cell noise."],
            ["Object", "Adds a custom_object_placement for a library object.", "Use for doors, crates, props, blockers, and interactive fixtures."],
            ["Interact", "Assigns dialogue to a placed object where supported.", "Use when an object should talk, open text, or lead to a scripted exchange."],
            ["Entity", "Places an entity_id on a cell with optional facing/schedule.", "Use for NPCs, enemies, companions, watchers, and test actors."],
            ["Trigger", "Places a step/interact/on_load/switch_change trigger tied to a cutscene.", "Use for scripted beats, tutorial gates, story flags, and systemic surprises."],
            ["Stamp", "Runs a preset builder at a clicked cell.", "Useful for repeatable rooms, corners, set-pieces, and generated structures."],
            ["Grid/Region controls", "Assigns region_id and emotional profile data.", "Use for Alderamontico Grid/Attend zones, survival pressure, and region-specific gates."],
        ],
        widths=[1.4, 2.0, 3.1],
    )
    add_para(
        doc,
        "The Inspector is the precise editor for map placements. Use it to add or edit entity placements, schedules, triggers, exits, world items, containers, region data, and conditions. The Lint overlay highlights validator warnings. The Play map button launches the current map directly for runtime verification.",
    )

    add_heading(doc, "6. Data Model Reference", 1)
    add_para(
        doc,
        "A game is a GamePackage. The following package fields are the main levers an AI uses to build content. Field names below match the current schema and should be preferred over invented aliases.",
    )
    add_table(
        doc,
        ["Package Field", "Contains", "Common AI Tasks"],
        [
            ["metadata", "title, version, start_map_id, start_spawn_id.", "Set the starting map and spawn for the build."],
            ["settings", "Loose settings such as player sprite, audio tracks, sound effects, and project-level knobs.", "Wire default player art, music, and presentation settings."],
            ["maps", "MapData records: cells, spawns, object placements, entity placements, items, containers, regions, triggers, exits.", "Build playable spaces and map graph."],
            ["object_library", "Object definitions with category, tags, tile_sprite_id, bounds, parts/mesh/asset, collision, simulation.", "Define walls, doors, props, blockers, workstations, and furniture."],
            ["sprite_library", "Sprite records with id, display name, dimensions, pixels or data_url.", "Register generated or pixel art for objects and actors."],
            ["entities", "Actor definitions with sprite, dialogue, stats, skills, emotional baselines, Attend nodes.", "Create NPCs, enemies, and party members."],
            ["dialogue", "Dialogue trees: nodes, options, conditions, switches, quest effects, cutscene triggers.", "Write conversations and choice consequences."],
            ["quests", "Quest objectives: talk, kill, collect, explore, interact, custom.", "Define the journal contract."],
            ["cutscenes", "Action sequences fired by triggers or dialogue options.", "Script state changes and scene flow."],
            ["switches", "Named boolean flags.", "Track story facts, one-time events, gates, and authored state."],
            ["items", "Inventory items with effects, category, sprite/icon, simulation, spatial profile.", "Make loot, keys, supplies, equipment, and consumables."],
            ["abilities", "Skills with AP/MP cost, element, targeting, range, payloads, emotional impulse.", "Build combat and systemic verbs."],
            ["shops", "Shop stock, prices, conditions, price modifiers.", "Build merchants and conditional economies."],
            ["barks", "Ambient lines between speaker pairs, gated by condition and cooldown.", "Make NPCs talk to each other and react to world state."],
            ["simulation_*", "Material profiles, processes, and workstations.", "Build material causality, crafting/process stations, traces, and tasks."],
            ["validators", "Validation configuration and extension data.", "Add project-specific audit contracts when useful."],
        ],
        widths=[1.45, 2.45, 2.6],
    )

    add_heading(doc, "7. Maps, Cells, Objects, and Exits", 1)
    add_para(
        doc,
        "MapData is the physical stage. It contains width, height, spawns, cells, props, custom_object_placements, entity_placements, item_placements, container_placements, regions, triggers, and exits. Every traversable slice should be checked for spawn-to-exit reachability.",
    )
    add_table(
        doc,
        ["Map Concept", "Important Fields", "How To Use It"],
        [
            ["Cell", "x, y, z, active, walkable, blocks_los, height, visual_height, terrain, object_id, region_id, room_id, tag, hazard, infection, portal_id, surface_tag, simulation.", "The base grid. Use active/walkable for traversal; blocks_los for visibility; surface_tag/simulation for chemistry seeds."],
            ["Object placement", "object_id, cell, facing, dialogue_id, blueprint_id.", "Places an object from the object library into the map."],
            ["Entity placement", "entity_id, cell, facing, schedule.", "Places an actor. Friendly NPC schedules can move by hour; hostiles use combat/perception AI."],
            ["World item placement", "id, item_id, cell, count.", "Loot on the ground. Runtime save deltas remember picked-up items."],
            ["Container placement", "id, object_id, cell, facing, display_name, locked, key_item_id, consume_key, items, simulation.", "Lootable blocker with persistent contents and optional lock/key behavior."],
            ["Region", "id, display_name, faction_id, reputation_threshold, survival_delta, passive_checks, alderamontico_grid, emotional_profile.", "Defines area pressure, access rules, emotional/Grid profiles, and survival effects."],
            ["Exit", "cell, target_map_id, target_spawn_id, facing, condition.", "Primary way to travel between maps. Use for stitched zones and doorways."],
        ],
        widths=[1.35, 2.55, 2.6],
    )
    add_callout(
        doc,
        "Closed doors and fog",
        "Closed doors and other closed blockers should count as line-of-sight blockers when fog or visibility is evaluated. If a door opens or closes at runtime, its map delta/state must be reflected in movement and LOS checks together.",
        fill="F2F7FC",
    )

    add_heading(doc, "8. Switches and Conditions", 1)
    add_para(
        doc,
        "Switches are named booleans stored in the package and runtime save. Use them for story facts, one-time triggers, gates, unlocked routes, seen scenes, consequences, and test toggles. A good switch name reads like a durable fact: watchfold_gate_open, girl_attended_once, reedmire_bridge_burned, or trial_marrowhouse_started.",
    )
    add_para(
        doc,
        "Conditions are declarative gates evaluated by the story service. They appear on dialogue options, triggers, cutscene branches, shops, price modifiers, exits, and barks. Predicates in one condition node are ANDed. The all, any, and not fields compose more complex logic.",
    )
    add_table(
        doc,
        ["Condition Field", "Meaning", "Example Use"],
        [
            ["switch / switch_value", "A switch must equal true or false.", "Show an option only after the player read a note."],
            ["quest / quest_state", "A quest is in a specific state.", "Open a cutscene branch after a trial starts."],
            ["has_item / item_count", "Inventory contains enough of an item.", "Allow unlocking a container with a key item."],
            ["party_contains", "An entity is currently in the party.", "Gate companion-specific dialogue."],
            ["faction with rep_gte / rep_lte", "Faction reputation is within bounds.", "Discount stock or bar access to a hostile region."],
            ["time_of_day", "Clock phase is late_night, night, dawn, day, or dusk.", "Swap night barks or nocturnal encounters."],
            ["hour_gte / hour_lt", "Hour window, including wraparound if start exceeds end.", "Run schedules or events between 22 and 5."],
            ["not / all / any", "Boolean composition.", "Express alternate solutions or negative gates."],
        ],
        widths=[1.7, 2.1, 2.7],
    )

    add_heading(doc, "9. Dialogue, Barks, Quests, and Documents", 1)
    add_para(
        doc,
        "Dialogue is for player-facing scenes and choices. A Dialogue has nodes; each node has a speaker, text, optional scene image, optional Attend data, and options. An option can move to another node, end the conversation, set switches, set quest state, trigger a cutscene, or be hidden by legacy required fields and/or a general condition.",
    )
    add_table(
        doc,
        ["Story Tool", "Best Use", "Avoid"],
        [
            ["Dialogue", "NPC conversations, choices, reactions, trial scenes, and object conversations.", "Huge lore speeches with no state change or choice."],
            ["Barks", "Ambient overheard exchanges between two entity ids when both speakers and the player are close enough.", "Replacing dialogue trees or quest logic."],
            ["Quests", "The player-facing goal contract and objective completion target ids.", "Using quests as the only record of tiny state; use switches for facts."],
            ["Documents", "Notes, books, terminal entries, letters, and readable lore.", "Putting mandatory logic only in prose without a switch or objective."],
        ],
        widths=[1.3, 2.7, 2.5],
    )
    add_para(
        doc,
        "For AI writing, pair each narrative scene with concrete state. A dialogue option should set a switch, update a quest, trigger a cutscene, start combat, open a shop, give/remove an item, or unlock a new condition. If nothing can change, make it a bark, document, or optional flavor node.",
    )

    add_heading(doc, "10. Events, Cutscenes, and Triggers", 1)
    add_para(
        doc,
        "Events are authored as Cutscene records. A trigger or dialogue option fires a cutscene; the cutscene runs a sequence of EventAction records. A cutscene can be blocking or nonblocking. Branch and label actions allow simple control flow without creating a new script language.",
    )
    add_table(
        doc,
        ["Trigger Type", "Fires When", "Typical Use"],
        [
            ["step", "The player enters the trigger cell.", "Ambushes, tutorials, discovery beats, first-visit flags."],
            ["interact", "The player acts on the trigger cell.", "Plaques, levers, doors, story objects, map fixtures."],
            ["on_load", "The map loads and conditions pass.", "Intro scenes, map-state cleanup, one-time arrivals."],
            ["switch_change", "A switch changes and conditions pass.", "Consequences that happen after dialogue, combat, or remote actions."],
        ],
        widths=[1.35, 2.2, 2.95],
    )
    add_table(
        doc,
        ["Cutscene Action Family", "Action Types"],
        [
            ["Movement and scene", "move_player, move_entity, teleport_player, camera_pan, screen_fade, wait."],
            ["Story and UI", "show_dialogue, read_document, open_shop, open_save_menu, game_end."],
            ["State", "set_switch, adjust_faction_rep, set_entity_hidden, branch, label, custom."],
            ["Rewards and inventory", "give_item, remove_item, give_currency, remove_currency, learn_skill."],
            ["Party and stats", "add_party_member, remove_party_member, heal_player, restore_party, modify_player_stats."],
            ["Combat and audio", "start_combat, play_sound, play_music."],
        ],
        widths=[2.0, 4.5],
    )
    add_callout(
        doc,
        "Event design pattern",
        "Use trigger -> cutscene -> switch as the basic chain. Example: stepping on a shrine trigger fires a cutscene; the cutscene shows dialogue, gives an item, sets shrine_seen true, and branches to a different line if the player already carries the matching relic.",
        fill="EAF4FF",
    )

    add_heading(doc, "11. Items, Shops, Skills, and Combat", 1)
    add_para(
        doc,
        "Items can be consumables, weapons, armor, or keys. They may restore health, magic, energy, hunger, thirst, fatigue, exposure, or add stat bonuses. Items can also carry simulation and spatial profiles, allowing them to matter to material systems and inventory packing.",
    )
    add_para(
        doc,
        "Shops expose stock with base prices, optional item conditions, and price modifiers. This lets AI create reputation shops, time-limited stock, quest-unlocked goods, and consequences without changing shop runtime code.",
    )
    add_para(
        doc,
        "Skills define AP cost, MP cost, element, targeting shape, range, payloads, and optional emotional impulse. Targeting shapes include single, line, cone, cross, and block. Elements include none, fire, shock, water, cold, poison, and physical. Payloads can damage, heal, apply status, or summon an entity.",
    )
    add_table(
        doc,
        ["Combat/System Surface", "Current Use"],
        [
            ["Turn/combat mode", "Same-map tactical combat with initiative, movement, attacks, skill use, enemy turns, and combat UI."],
            ["Perception", "Enemies can alert, search, investigate, report, or chase based on visibility, light, sound, gas, fire, and player presence."],
            ["Cover/flank/overwatch", "Stage 6 combat facts and overlays support shove, overwatch, cover edges, height/facing modifiers, and telegraphed intent."],
            ["Chemistry bridge", "Fire, water, cold, shock, poison, foam, oil, steam, smoke, gas, ice, wetness, and physical states can alter cells and actors."],
            ["Emotional skills", "Skills or built-in verbs can push Alderamontico axes such as valence, arousal, grief, reverence, and attachment."],
        ],
        widths=[1.8, 4.7],
    )

    add_heading(doc, "12. Simulation, Chemistry, Grid, and Attend", 1)
    add_para(
        doc,
        "The engine now has a layered systemic substrate. Simulation materials describe density, hardness, flammability, absorbency, conductivity, fragility, scent retention, cleaning difficulty, decay, sound, and light response. Authored simulation profiles can live on cells, objects, items, and containers. Processes and workstations support contextual multi-step interactions such as starting, working, canceling, and collecting outputs.",
    )
    add_para(
        doc,
        "Chemistry uses numeric axes and derived conditions. Elemental verbs and surface tags can create fire, wetness, ice, steam, smoke, toxic gas, conduction, smothering, scorching, melting, and actor statuses. When AI adds a systemic set-piece, it should declare which shipped verbs solve it and verify that player feedback exists on the tile, actor, log, HUD, audio, or event surface.",
    )
    add_para(
        doc,
        "The Alderamontico layer adds actor emotional axes, Attend read-outs, Grid amplification by region, Glass residue, emotional profiles, lens actors, emotional skills, and behavior modes. This layer should remain integrated with the normal systems: it is not a second dialogue system, not a global emotional field, and not a replacement for switches or quests.",
    )
    add_table(
        doc,
        ["System", "Authoring Hook", "Player Feedback Requirement"],
        [
            ["Material simulation", "simulation_materials plus simulation profiles on cells/objects/items.", "Visible condition, trace, process, or object state changes."],
            ["Chemistry", "surface_tag, elemental skills, global verbs, authored seed surfaces.", "Screen/log/HUD/world feedback for on fire, wet, freezing, poisoned, etc."],
            ["Workstations/processes", "simulation_workstations and simulation_processes.", "Clear prompt for Start, Work, Cancel, Collect, and costs."],
            ["Perception/stealth", "Light/sound/visibility stimuli, hostile AI, region/map layout.", "Alert state, sight lines, search behavior, barks, and logs should move with the actor."],
            ["Grid/Attend", "Entity attend_node, emotional_axes, region emotional_profile, alderamontico_grid.", "Attend panel, surface-vs-hidden read, pressure, Glass, Grid chip, and log/audio feedback."],
        ],
        widths=[1.55, 2.55, 2.4],
    )

    add_heading(doc, "13. Asset and Overworld Production", 1)
    add_para(
        doc,
        "For the current overworld, the active art standard is generated bitmap art with square-faced oblique terrain, structure, barrier, and prop sprites: front and top faces visible, sketchy painted/Da Vinci notebook texture, no labels, no baked backgrounds, no isometric taper, and no single-tile town/city/fracture symbols. Tiles must be cropped out of atlas gutters and edge-conditioned to loop. Props should be transparent cutouts in the same perspective.",
    )
    add_table(
        doc,
        ["Asset Script", "Purpose"],
        [
            ["npm run audit:overworld-assets", "Validates fallback/source pixel library and writes palette, manifest, style reference, and contact sheet."],
            ["npm run art:extract-oblique-tiles", "Crops terrain atlas, removes gutters, edge-conditions looping tiles, writes manifest/contact sheet."],
            ["npm run art:extract-oblique-structures", "Crops square-faced wall/door structure tiles and rejects angled/isometric attempts."],
            ["npm run art:extract-oblique-barriers", "Crops barrier/aperture sprites with transparency and seam handling."],
            ["npm run art:extract-oblique-props", "Crops prop atlases, chroma-keys backgrounds to alpha, writes prop manifests/contact sheets."],
            ["npm run art:extract-player", "Crops and pads Intercessor directional idle/step player frames."],
        ],
        widths=[2.3, 4.2],
    )
    add_para(
        doc,
        "Doc 07 currently has Phase 0-1 art and Phase 2-3 geography/greybox grounded. Phase 4 population is the next overworld target: systemic set-pieces, sparse ambient fill, soft-gates, discoveries, enemies, loot, and sidequest hooks across the nine March maps.",
    )

    add_heading(doc, "14. AI Authoring Checklists", 1)
    add_heading(doc, "New Map Checklist", 2)
    add_bullets(
        doc,
        [
            "Define map id, display name, size, role, and start/edge spawns.",
            "Paint active/walkable cells with broad readable terrain patches.",
            "Place blocked/LOS cells, walls, closed doors, void, and barriers intentionally.",
            "Add exits to connected maps and verify target spawns/facing.",
            "Add objects, containers, items, entity placements, regions, and triggers.",
            "Run npm run audit:maps and playtest start-to-exit traversal.",
        ],
    )
    add_heading(doc, "New Quest or Scene Checklist", 2)
    add_bullets(
        doc,
        [
            "Create a quest only if the player needs journal tracking; otherwise use switches and dialogue.",
            "Name switches as durable facts, not temporary UI labels.",
            "Build dialogue options with conditions, set_switch, trigger_quest, or trigger_cutscene as needed.",
            "Author a cutscene for multi-action consequences.",
            "Use triggers for place-based or switch-change consequences.",
            "Test both the first-time and repeated paths.",
        ],
    )
    add_heading(doc, "New NPC or Enemy Checklist", 2)
    add_bullets(
        doc,
        [
            "Create an entity with sprite_id, stats, dialogue_id or party_dialogue_id, skills, and is_npc when appropriate.",
            "Add emotional_axes or attend_node only when the content needs them.",
            "Place the entity with facing and optional schedule.",
            "If hostile, test sight, chase, search, combat start, and post-attack pursuit.",
            "If ambient, add barks with speaker pair, lines, condition, and cooldown.",
        ],
    )
    add_heading(doc, "New Systemic Set-piece Checklist", 2)
    add_bullets(
        doc,
        [
            "Declare the problem in map terms: blocked route, hazard, guarded object, soft gate, or discovery.",
            "Use at least two shipped systems: chemistry, movement, cover, perception, containers, doors, skills, Grid, or inventory.",
            "Provide at least two solutions expressible in existing verbs such as push, pull, throw, burn, douse, freeze, wet, electrify, foam, break, climb, or unlock.",
            "Make feedback visible in the world and HUD/log/audio where relevant.",
            "Add an auditable note, marker, metadata field, or validator hook if the set-piece belongs to overworld Phase 4.",
            "Playtest the obvious solution, one alternate solution, and failure/repeat behavior.",
        ],
    )
    add_heading(doc, "Release or Handoff Checklist", 2)
    add_bullets(
        doc,
        [
            "Run npm run lint and npm run build for code-facing changes.",
            "Run targeted tests: test:engine, test:chemistry, or test:state when touched.",
            "Run audits: audit:maps, audit:combat, audit:overworld, audit:overworld-assets as relevant.",
            "Use browser QA for map visuals, Play mode, UI overlays, generated sprites, and user-reported bugs.",
            "Update implementation docs, manifests, and phase status when the delivered slice changes project state.",
        ],
    )

    add_heading(doc, "15. Common Patterns", 1)
    add_table(
        doc,
        ["Goal", "Recommended Chain"],
        [
            ["Locked door with key", "Item key -> door/container condition -> interact trigger or object state -> set door_open switch -> movement/LOS update."],
            ["One-time discovery", "Step trigger once -> cutscene read_document/show_dialogue -> set discovered switch -> quest explore objective updates."],
            ["Conditional merchant", "Shop item condition by quest/switch/faction -> price_modifiers by reputation -> dialogue option open_shop."],
            ["Ambient town reaction", "Two NPC placements with barks -> bark condition by switch/time -> cooldown_minutes."],
            ["Soft-gated hazard", "Blocked or hazardous map route -> chemistry/verb solution -> switch or map delta opens shortcut -> audit ensures bypass exists."],
            ["Story trial beat", "Dialogue scene -> switch and quest state -> cutscene action -> region/emotional profile or Attend node consequence."],
            ["Combat tutorial", "Step trigger -> start_combat -> enemy placement with clear sight/perception -> skill/overwatch feedback -> quest objective."],
        ],
        widths=[1.7, 4.8],
    )

    add_heading(doc, "16. Glossary", 1)
    add_table(
        doc,
        ["Term", "Meaning"],
        [
            ["Package", "The authored game data loaded by the Studio and Play runtime."],
            ["Save delta", "Runtime changes stored separately from authored map data: opened doors, looted items, moved objects, fog, chemistry, and state."],
            ["Switch", "A named boolean fact used for gating and consequences."],
            ["Condition", "A declarative predicate over switches, quests, items, party, faction reputation, and time."],
            ["Cutscene", "A sequence of EventAction records, not necessarily a cinematic."],
            ["Trigger", "A map or state hook that fires a cutscene."],
            ["Bark", "Short ambient NPC-to-NPC exchange, gated by proximity, condition, and cooldown."],
            ["Object library", "Reusable definitions for placed objects, walls, doors, props, blockers, and interactables."],
            ["Sprite library", "Image records used by players, entities, objects, tiles, and generated art."],
            ["Region", "A named map area with survival, faction, passive check, Grid, or emotional profile data."],
            ["Attend", "A player action/read-out for authored emotional/hidden condition surfaces."],
            ["Grid", "The Alderamontico amplification layer operating through regions, lenses, emotional axes, and Glass."],
            ["Audit", "A script that rejects broken content before it reaches hand QA."],
        ],
        widths=[1.6, 4.9],
    )

    add_callout(
        doc,
        "Bottom line",
        "A strong AI-built slice in this engine is not just text, not just code, and not just art. It is a package change that can be authored in the Studio, played in Play mode, understood through visible feedback, persisted through saves, and checked by scripts.",
        fill="EAF4FF",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
